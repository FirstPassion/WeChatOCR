import wcocr
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from colorama import init, Fore, Style


def find_wechat_path():
    # 获取当前脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    # 构建 WeChatOCR 的路径,当前目录下的path文件夹是否存在
    common_paths = os.path.join(script_dir, "path")
    # 检查路径是否存在
    if os.path.exists(common_paths):
        return common_paths
    else:
        print(f"{common_paths} 文件夹不存在")
        return None


def find_wechatocr_exe():
    # 获取 WeChatOCR.exe 的完整路径
    script_dir = os.path.dirname(os.path.abspath(__file__))
    # 检查当前目录下的path文件夹下面的WeChatOCR文件夹已经目录下面的WeChatOCR.exe文件是否存在
    wechatocr_path = os.path.join(script_dir, "path", "WeChatOCR", "WeChatOCR.exe")
    # 检查文件是否存在
    if os.path.isfile(wechatocr_path):
        return wechatocr_path
    else:
        print(f"{wechatocr_path} 文件不存在")
        return None


def wechat_ocr(image_path):
    """
    微信OCR识别函数
    :param image_path: 图片路径
    :return: 识别结果
    """
    # 查找 WeChatOCR 的路径
    wechat_path = find_wechat_path()
    # 查找 WeChatOCR.exe 的路径
    wechatocr_path = find_wechatocr_exe()
    # 检查路径是否存在
    if not wechat_path or not wechatocr_path:
        return []  # 返回空结果

    # 初始化 WeChatOCR
    wcocr.init(wechatocr_path, wechat_path)
    # 执行 OCR 识别
    result = wcocr.ocr(image_path)
    texts = []

    # 处理 OCR 结果
    for temp in result["ocr_response"]:
        text = temp["text"]
        # 检查text是否是bytes类型
        if isinstance(text, bytes):
            # 解码字节为字符串
            text = text.decode("utf-8", errors="ignore")
        texts.append(text)
    # 返回识别结果
    return texts


def save_to_docx(texts, output_path):
    """
    保存识别结果到docx文件
    :param texts: 识别结果
    :param output_path: 输出路径
    """
    # 创建一个新的 Word 文档
    doc = Document()
    # 遍历识别结果
    for text in texts:
        # 添加段落并设置宋体字体
        paragraph = doc.add_paragraph()
        # 添加段落内容
        run = paragraph.add_run(text)
        # 设置字体为宋体
        run.font.name = "宋体"
        # 设置字体为宋体 (兼容中文设置)
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 设置字体大小为五号字体 (10.5 磅)
        run.font.size = Pt(10.5)

    # 保存文档到指定路径
    doc.save(output_path)


def process_all_images():
    """
    处理所有图像
    """
    # 获取当前脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    src_folder = os.path.join(script_dir, "src")  # 源文件夹路径
    docx_folder = os.path.join(script_dir, "docx")  # 输出文档夹路径
    # 如果输出文件夹不存在，则创建
    os.makedirs(docx_folder, exist_ok=True)
    # 支持的图像格式
    image_extensions = (".jpg", ".jpeg", ".png", ".bmp", ".tif")
    # 遍历 src 文件夹及其所有子文件夹
    for root, dirs, files in os.walk(src_folder):
        for file in files:
            if file.lower().endswith(image_extensions):
                image_path = os.path.join(root, file)  # 获取图像文件的完整路径
                relative_path = os.path.relpath(root, src_folder)  # 获取相对路径
                docx_folder_path = os.path.join(
                    docx_folder, relative_path
                )  # 输出文档的文件夹路径

                # 确保 docx 文件夹路径存在
                os.makedirs(docx_folder_path, exist_ok=True)

                # 处理图片文件
                print(
                    Fore.GREEN
                    + f"正在处理: {os.path.relpath(image_path, script_dir)}"
                    + Style.RESET_ALL
                )
                texts = wechat_ocr(image_path)  # 调用 OCR 函数
                image_name = os.path.splitext(file)[0]  # 获取图像文件名
                output_docx = os.path.join(
                    docx_folder_path, f"{image_name}_OCR.docx"
                )  # 输出文档路径
                save_to_docx(texts, output_docx)  # 保存 OCR 结果到文档
                # 显示相对路径
                relative_docx_path = os.path.relpath(output_docx, script_dir)
                print(f"OCR 结果已保存到： {relative_docx_path}\n")


if __name__ == "__main__":
    init(autoreset=True)  # 初始化 colorama
    process_all_images()  # 处理所有图像
    print(Fore.RED + "全部文件处理完成，请按 Enter 键退出……" + Style.RESET_ALL)
    input()  # 等待用户输入
